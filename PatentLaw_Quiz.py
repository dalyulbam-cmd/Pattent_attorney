import os, random, itertools, time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


##################################################################################

OrderInLaw = ["편","장","절","관","조","항","호","목"]
paragraph_symbol = ['①','②','③','④','⑤','⑥','⑦','⑧','⑨','⑩','⑪','⑫','⑬','⑭','⑮','⑯','⑰','⑱','⑲','⑳']
subparagraph_symbol = ['1.','2.','3.','4.','5.','6.','7.','8.','9.','10.']
item_symbol = ['가','나','다','라','마','바','사','아','자','차','카','타','파','하']
#common은 단순한 접속사 및 말의 어미로, 예측이 지극히 쉬운 경우를 다룬다.
#"대한", "관련한"등은 단순하고 반복적이지만, 조문의 구조를 이해하는데 도움이 되므로 포함하지 않는다. 
common = ["및","또는","그","있다.","있다.\n","한다.","한다.\n","본다.","본다.\n",
          "각","호의","목의","아니한다.","아니한다.\n","할","수","것","아니한"
          ,"이하","없다.","없다.\n","등","경우","어느"]

"""
https://www.lawmaking.go.kr/lmKnlg/jdgStd/info?astSeq=96&astClsCd=
대한민국의 법령은 위와 같이 Index를 나눈다. 
Article은 조, Paragraph는 항, Subparagraph는 호, Item은 목을 영어로 표기한 것이다.  
"""

##################################################################################

"""
templit은 읽어들일 조문이 들어있는 워드 파일.
Result는 출제된 문제가 들어있는 워드 파일.
NP는 number of problem, 즉 문제의 개수.
Keyword는 출제하고 싶은 문제와 관련된 키워드.

파일 리스트 :
    특허 :
        특허Total.docx 
        특허총칙.docx
        특허요건.docx
        특허심사.docx
        특허등록.docx
        특허권.docx
        특허권자의보호.docx
        특허취소신청.docx
        특허심판.docx
        특허재심.docx
        특허소송.docx
        특허국제출원.docx
        특허보칙.docx
        특허벌칙.docx
    상표 :
        상표Total.docx
        상표총칙.docx
        상표요건.docx
        상표심사.docx
        상표등록.docx
        상표권.docx
        상표권자의보호.docx
        상표재심소송.docx
        상표상품분류전환등록.docx
        상표국제출원.docx
        상표보칙.docx
        상표벌칙.docx
    디자인 :
    민법 :
        민법Total.docx
        민법총칙법인.docx
        민법총칙주소.docx
        민법총칙능력.docx
        민법총칙기간과소멸시효.docx
        민법총칙물건과법률행위.docx
        
    민사소송법 : 
        
        
"""
# NP = Number of Problem, NB = Number of Blank, LB = Length of Blank
# 문제 수를 조절하세요. 
NP = 3
NB = 4
LB = 2

#현재 프로그램의 디렉토리를 확인하세요. 
Current = os.getcwd()
print(Current)

# 원하는 법조문 Templit을 가져와 문제를 만드세요. 
templit = Document('자료/특허심판.docx')
Vacant = Document('Today.docx')
Result = "Today's_Problem_" + str(time.strftime("%y-%m-%d-%H-%M")) +".docx"

Your_Keyword = None 


##################################################################################

class Problem(object):
    
    def __init__(self):

        self.Whole_text = []
        self.Problem_text = ""
        self.Last_class_name = ["","","",""]
        self.Last_class = 0
        self.Keyword = []

    def put_data(self,Line_Type,Text):
        """
        데이터 입력함수. 
        조, 항, 호, 목의 본문은 각각 하나의 Line에 할당된다.
        즉, Line_Type은 "제5조의2", "제1항" 같은 Str 타입의 정보를 담고 있다.
        다만, 하나의 Line에 조,항이 같이 포함되는 등, 복잡한 경우가 있어 주의해야한다.
        """ 
        (val_A, val_P, val_S, val_I) = Line_Type

        # 문제 class에서는 조문을 구분하는 가장 높은 단위는 '조'다. 
        if val_A:
            self.Last_class_name[0]= val_A

        # 제1항 다음에 나오는 본문이 제2항이라면, present class는 1, last class도 1. 
        # 제1호 다목 다음에 나오는 본문이 2호라면, present class는 2, last class는 3. 
        for i in range(len(Line_Type)):
            if Line_Type[i]:
                present_class = i

        #제1항제1호다목 -> 제1항제2호 ;
        #["제10조","제1항","제1호",다목] -> ["제10조","제1항","제2호",""]
        self.Last_class_name[present_class] = Line_Type[present_class]
        if present_class <= self.Last_class:
            for i in range(len(self.Last_class_name)):
                if i > present_class:
                    self.Last_class_name[i] = ""

        total_class_name = ""
        for name in self.Last_class_name:
            total_class_name += str(name)
        self.Whole_text.append([total_class_name,Text])        
        self.Last_class = present_class 
    
    def blank(self,NB,LB):
        """
        한 조항의 모든 항,호,목의 텍스트를 일률적으로 합친다.
        모드에 따라 빈칸을 내는 방법이 다르나, 재조립하는 방법은 같다.
        """
        self.Problem_text = All_Text(self.Whole_text)
        Problem_text_list1 = [word.strip(" ") for word in self.Problem_text.split(" ")]
        # 조문에 LB만큼 연속된 빈칸을 NB개 랜덤하게 만든다. 
        Problem_text_list2 = random_blank_wiselong(Problem_text_list1,NB,LB)
        #빈칸을 넣은 텍스트를 재조립  
        self.Problem_text = " ".join(Problem_text_list2)
        return self.Problem_text

    def check_trash(self):
        data = self.Whole_text[0][1]
        data1 = data.split(" ")
        data2 = [word.strip(" ") for word in data1 if word.strip(" ")]
        if data2[1] == "삭제":
            return 1
        else :
            return 0

##################################################################################

def All_Text(The_list):
    All = "\n"
    for text in The_list:
        All += str(text[1]) + "\n"
    return All

def All_Text2(The_list):
    All = "\n"
    for text in The_list:
        All += str(text) + "\n"
    return All

#################################################################################

def random_long_sublist(the_list,NB,LB):
    order_list = []
    sublist = []
    end = len(the_list)-1 
    start = LB -1 
    def drawing(order_list):
        n = random.randint(start,end)
        group = []
        for i in range(LB):
            value = n-i
            if value in order_list:
                return drawing(order_list)
            group.append(n-i)
        order_list += group
        return order_list
    for i in range(NB):
        order_list = drawing(order_list)
    for j in order_list:
        sublist.append(the_list[j])
    return order_list, sublist


def random_long_blank(the_list,NB,LB):
    order_list, sublist = random_long_sublist(the_list,NB,LB)
    for i in range(len(sublist)):
        #뻔한 단어가 Blank되었을 때, '뻔하지 않을 때까지' 랜덤으로 하나를 뽑는다.
        while sublist[i] in (common + item_symbol + paragraph_symbol + subparagraph_symbol + [" ","\n"]):
            little_order_list, little_sublist = random_long_sublist(the_list,1,1)
            #이미 뽑힌 적 있는 단어라면, 다시 뽑는다. 
            if little_order_list[0] in order_list:
                continue
            order_list[i] = little_order_list[0]
            sublist[i] = little_sublist[0]
        word = sublist[i]
        blank = ""
        for j in range(len(word)):
            #글자 수를 알 수 없게 하는 장치. 
            if j%4 ==0 : 
                blank += "    "
        blank = "(" + blank + ")"
        sublist[i] = blank
    #Synchronizing 
    for i in range(len(sublist)):
        the_list[order_list[i]] = sublist[i]
    return the_list    


def random_blank_wiselong(the_list,NB,LB):
    trim = []
    for word in the_list:
        if '"' in word :
            trim.append(the_list.index(word))
    if trim == []:
        return random_long_blank(the_list,NB,LB)
    else :
        edited_list = the_list[trim[0]:trim[1]+1]
        return random_long_blank(edited_list,NB,LB)

#################################################################################

def replaceRight(original, old, new, count_right):
    """
    https://yuddomack.tistory.com/entry/파이썬-replace-문자열-제거-수정변환
    ...에서 함수의 내용을 참조.
    Str class에서 replace 함수가 '좌측우선인것'과 '우측우선함수'가 없다는 문제가 있었다.
    다른 제작자분이 직접 만들어 놓으셨다. 검색해서 표출이 되지 않았다면 직접 만들 심산이었으나
    어깨가 가벼워졌다. 
    """
    repeat = 0
    text = original
    old_len = len(old)

    count_find = original.count(old)
    if count_right > count_find :
        repeat = count_find
    else :
        repeat = count_right
    while(repeat):
        find_index = text.rfind(old)
        text = text[:find_index] + new + text[find_index+old_len:]
        repeat -= 1
    return text 


##################################################################################

def Article_finder(text_list):
    # ['제3조',...]와 같이 구분된 텍스트에서 '제'를 찾고, 제3조의2 같은 신설조항을 구분해준다. 
    text = text_list[0]
    total = ""
    if text[0] == '제':
        for letter in text:
            try :
                #type(int(0))== int => True , int(0) or False => False 
                if (letter == "의" or letter == OrderInLaw[4] or type(int(letter))==int):
                    total += letter
                if len(total)>2 and type(int(total[1]))==int and (total[-1] in OrderInLaw[0:3]):
                    return 0
            except :
                continue
        while total[-1] == "의":
            total = replaceRight(total,'의','',1)
        return "제" + total 
    else :
        return 0 


def Paragraph_finder(text_list):
    # 텍스트에서 처음 표기된 특수문자를 찾고 값을 되돌린다.(예: '①','②'...)
    total = ""
    for unit in text_list:
        if unit in paragraph_symbol:
            total = unit
            break 
        else :
            continue
    if total == "":
        return 0
    else :
        return '제' + unit + '항'
    

def Subparagraph_finder(text_list):
    # '1.' 같이 적힌 호 표기를 찾는다. 온점을 찾은 다음, int 값으로 변환가능한지 본다.  
    total = ""
    text = text_list[0]
    total = text[:-1]
    if text[-1] == "." :
        try :
            int(text[:-1])
        except :
            return 0
        return "제" + total + "호"
    else :
        return 0 

def Item_finder(text_list):
    # '가.' 같이 적힌 목 표기를 찾는다. 온점을 찾은 다음, 가나다라...중 하나인지 확인한다. 
    total = ""
    text = text_list[0]
    total = text[:-1]
    if text[-1] == "." and (total in item_symbol):
        return total + "목"
    else :
        return 0 


##################################################################################

"""
https://m.blog.naver.com/anakt/221842622079
다음 조항으로 넘어갈 때 '문제 유닛'을 클래스 선언하고, 이전 '문제 유닛'을 '문제 리스트'에 넣는다.
"""
def data_acquisition(templit):
    Problem_list = []
    Problem_unit = None
    
    for row, paragraph in enumerate(templit.paragraphs):
        if row > 100:
            break
        """
        첫 째로 라인을 띄어쓰기 가능한 단위로 쪼개고,
        둘 째로 각 단어에서 공백을 제거하고,
        셋 째로 \n을 "\n" + " "로 만든다. 
        """
        line1 = paragraph.text.split(' ')
        line2 = [word.strip(" ") for word in line1 if word.strip(" ")]
        if line2 == []:
            continue
        Arranged_text = " ".join(line2) + "\n" + " "
        
        Line_Type = (val_A, val_P, val_S, val_I) = (Article_finder(line2), Paragraph_finder(line2), Subparagraph_finder(line2), Item_finder(line2)) 
        if Line_Type == (0,0,0,0):
            continue
        
        if val_A :
            if Problem_unit:
                if not Problem_unit.check_trash():
                    Problem_list.append(Problem_unit)
            Problem_unit = Problem()
        Problem_unit.put_data(Line_Type,Arranged_text)
    Problem_list.append(Problem_unit)
    return Problem_list

##################################################################################    
"""
여기 함수들은 출제방식을 나타낸다.
Set_Exam_Random()는 리스트에 있는 조문을 랜덤하게 뽑기만 한다.
NP는 Number of Problem의 약자, NB는 Number of Blank의 약자다. 
"""

def Set_Exam_AllRandom(Vacant,Problem_list,NP,NB,LB):

    for i in range(NP):
        new_problem = random.choice(Problem_list)
        Problem_list.remove(new_problem)
        new_problem.blank(NB,LB)
        Vacant.add_paragraph(new_problem.Problem_text)
    Vacant.save(Result)

##################################################################################
        
Problem_list = data_acquisition(templit)
Set_Exam_AllRandom(Vacant,Problem_list,NP,NB,LB)




    
    
