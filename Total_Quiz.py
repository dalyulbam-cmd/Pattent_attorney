import os, random, itertools, time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

################################################################################
################################################################################
################################################################################

# 원하는 법조문 Templit을 가져와 문제를 만드세요. 
Current = os.getcwd()
print(Current)
templit = Document('문제은행-민법/민법물권각칙판례.docx')
Vacant = Document('Today.docx')
Result = "Today's_Problem_" + str(time.strftime("%y-%m-%d-%H-%M-%S")) +".docx"

keyword_list = []

# NP = Number of Problem, NB = Number of Blank, LB = Length of Blank
# 문제 수를 조절하세요. 
NP,NB,LB = 3,2,3


################################################################################
################################################################################
################################################################################
"""
☆리스트

    [판례집]
    
        <특허>
            특허심판판례.docx
            특허정의판례.docx
            특허요건판례.docx
        <상표>
        <민법>
            민법총칙법인판례.docx

    [조문집]
    
        <특허>
            특허Total.docx 
            특허총칙.docx
            특허요건.docx
            특허심사.docx
            특허등록.docx
            특허권.docx
            특허권자의보호.docx
            특허취소신청.docx
            특허심판.docx
            특허재심.docx
            특허소송.docx
            특허국제출원.docx
            특허보칙.docx
            특허벌칙.docx
        <상표>
            상표Total.docx
            상표총칙.docx
            상표요건.docx
            상표심사.docx
            상표등록.docx
            상표권.docx
            상표권자의보호.docx
            상표재심소송.docx
            상표상품분류전환등록.docx
            상표국제출원.docx
            상표보칙.docx
            상표벌칙.docx
        <디자인>
            디자인Total.docx
            디자인총칙.docx
            디자인등록출원.docx
            디자인심사.docx
            디자인등록.docx
            디자인권.docx
            디자인권보호.docx
            디자인심판.docx
            디자인재심소송.docx
            디자인국제출원.docx
            디자인보칙.docx
            디자인벌칙.docx
        <민법>
            민법Total.docx
            민법총칙법인.docx
            민법총칙주소.docx
            민법총칙능력.docx
            민법총칙기간과소멸시효.docx
            민법총칙물건과법률행위.docx
        <민사소송법>
        <타법>
            발명진흥법.docx
        <시행령>
            특허시행령.docx
        <시행규칙>
            특허시행규칙.docx

    [테마]
    
        <특허>
            그밖의발명.docx
            미생물발명.docx
            방법의사용을청약하는행위.docx
            특허이용관계.docx
            
            특허에관한절차일반.docx
            특허출원서류.docx
            특허절차보정.docx
            특허받을수있는대상.docx
            특허권리행위능력.docx
            특받권특허권공유.docx
            특허제출기간유예.docx
            직무발명.docx
            
            특허요건조문모음.docx
            특허실체보정.docx

            특허이익제도.docx

            


"""
################################################################################
################################################################################
################################################################################


OrderInLaw = ["편","장","절","관","조","항","호","목"]
paragraph_symbol = ['①','②','③','④','⑤','⑥','⑦','⑧','⑨','⑩','⑪','⑫','⑬','⑭','⑮','⑯','⑰','⑱','⑲','⑳']
subparagraph_symbol = ['1.','2.','3.','4.','5.','6.','7.','8.','9.','10.','11.','12.','13.','14.','15.','16.','17.','18.','19.','20.']
item_symbol = ['가','나','다','라','마','바','사','아','자','차','카','타','파','하','거']
#common은 단순한 접속사 및 말의 어미로, 예측이 지극히 쉬운 경우를 다룬다.
#"대한", "관련한"등은 단순하고 반복적이지만, 조문의 구조를 이해하는데 도움이 되므로 포함하지 않는다. 
common = ["있다","한다","본다","아니다","없다","아니한다","판시하였다",
          "각","그","및","할","수","것","이","볼","등",
          "또는","이는","이를","이러한","아니한",
          "것인","것은","것으로서","것이고","것을","것이",
          "호의","목의",
          "경우","어느","따라","여하","이하",
          "뿐만","라고",
          "여기서",
          "대한","대하여","의한","의하여","관한","관하여"]
Common_and_Symbol = common + item_symbol + subparagraph_symbol + paragraph_symbol + OrderInLaw

"""
https://www.lawmaking.go.kr/lmKnlg/jdgStd/info?astSeq=96&astClsCd=
대한민국의 법령은 위와 같이 Index를 나눈다. 
Article은 조, Paragraph는 항, Subparagraph는 호, Item은 목을 영어로 표기한 것이다.  
"""

################################################################################
################################################################################
################################################################################

def All_Text(The_list):
    All = "\n"
    for text in The_list:
        All += str(text[1]) + "\n"
    return All

def All_Text2(The_list):
    All = "\n"
    for text in The_list:
        All += str(text) + "\n"
    return All

def random_distribution(NU,NB):
    # 합이 NB인 NU개 이하의 숫자들의 리스트를 리턴한다.
    count = 0
    NL = []
    while NB > 0 and count < NU:
        n = random.randint(1,NB)
        NB -= n
        count += 1
        NL.append(n)
    # Give last element all Left Number from NB
    NL[-1] += NB
    random.shuffle(NL)
    return NL


def replaceRight(original, old, new, count_right):
    """
    https://yuddomack.tistory.com/entry/파이썬-replace-문자열-제거-수정변환
    ...에서 함수의 내용을 참조.
    Str class에서 replace 함수가 '좌측우선인것'과 '우측우선함수'가 없다는 문제가 있었다.
    다른 제작자분이 직접 만들어 놓으셨다. 검색해서 표출이 되지 않았다면 직접 만들 심산이었으나
    어깨가 가벼워졌다. 
    """
    repeat = 0
    text = original
    old_len = len(old)

    count_find = original.count(old)
    if count_right > count_find :
        repeat = count_find
    else :
        repeat = count_right
    while(repeat):
        find_index = text.rfind(old)
        text = text[:find_index] + new + text[find_index+old_len:]
        repeat -= 1
    return text 

def range_involved(number,exception):
    confirmation = False
    if exception == []:
        return confirmation
    while not exception == []:
        if exception[0] <= number and number <= exception[1]:
            confirmation = True
            break
        else :
            exception.pop(0)
            exception.pop(0)
    return confirmation

################################################################################
################################################################################
################################################################################

def text_involved(text,the_list):
    """
    the_list = ['<','>'] 라고 했을 때,
    text = "<개정2014.4.16." 이라면, '<' in text == True 가 성립한다.
    """
    confirmation = False
    for i in range(len(the_list)):
        if the_list[i] in text :
            confirmation = True
    return confirmation

def text_involved2(text,the_list):
    """
    the_list = ['될','돼']  라고 했을 때,
    text = "될까?" 이라면, '될' in text == True 가 성립한다.
    그러나 text = "북한산돼지" 이라면, '돼' in text == False 가 성립한다. 
    """
    confirmation = False
    for i in range(len(the_list)):
        if the_list[i] in text and not len(text) > len(the_list[i]) + 2 :
            confirmation = True
    return confirmation 
        

################################################################################
################################################################################
################################################################################


def random_long_sublist(the_list,NB,LB):
    order_list = []
    sublist = []
    end = len(the_list)-1 
    start = LB -1 
    
    def drawing(order_list):
        n = random.randint(start,end)
        group = []
        for i in range(LB):
            value = n-i
            if value in order_list:
                return drawing(order_list)
            group.append(n-i)
        order_list += group
        return order_list
    for i in range(NB):
        order_list = drawing(order_list)
    for j in order_list:
        sublist.append(the_list[j])
    return order_list, sublist


def random_long_blank(the_list,NB,LB,exception):
    order_list, sublist = random_long_sublist(the_list,NB,LB)
    """
    만약 무작위로 뽑아낸 리스트가 우연히도 exception 안에 들어있거나,
    평범하거나 심볼에 불과한 경우에는 다시 뽑도록 설계했다.
    각각 check1 과 check2로서 해당 여부는 bool 값으로 표시된다. 
    """
    for i in range(len(sublist)):
        #check1 = text_involved2(sublist[i],Common_and_Symbol)
        #check2 = range_involved(order_list[i],exception) -> 다음에 implement
        while sublist[i] in (common + item_symbol + paragraph_symbol + subparagraph_symbol + [" ","\n"]):
            little_order_list, little_sublist = random_long_sublist(the_list,1,1)
            if little_order_list[0] in order_list:
                continue
            order_list[i] = little_order_list[0]
            sublist[i] = little_sublist[0]

        word = sublist[i]
        blank = ""
        for j in range(len(word)):
            #글자 수를 알 수 없게 하는 장치. 
            if j%4 ==0 : 
                blank += "    "
        blank = "(" + blank + ")"
        sublist[i] = blank
        
    for i in range(len(sublist)):
        the_list[order_list[i]] = sublist[i]
    return the_list    

def random_point_blank(the_list,NB,LB):
    trim = []
    exception = []
    """
    trim은 해당 파트만 blank로 대체하는 편집을 하겠다는 표시이고,
    exception은 해당 파트를 blank 범위에서 제외하고 편집을 하겠다는 표시다. 

    큰 따옴표로 둘러쌓인 문장이 있다면, 중요표시로 보고, trimming target으로 놓는다. 
    이와 달리 < > 화살표 괄호로 둘러쌓인 문장이 있다면, 판례번호 인지 개정연혁에 관한 표시인지 확인할 필요가 있다. 
    따라서 < > 괄호의 내용을 target exception에 집어넣는다. exception은 the_list의 index를 원소로 하는 리스트다. 
    """
    for word in the_list:
        if '"' in word :
            trim.append(the_list.index(word))
        """
        if text_involved(word,['<','>','[',']']):
            exception.append(the_list.index(word))
            if exception.count(the_list.index(word))>2:
                exception.pop(-1)
                print(exception)
        """
    if trim == []:
        return random_long_blank(the_list,NB,LB,exception)
    else :
        #LN is a shorten word for Left Number
        #NU is a shorten word for Number of Unit
        #NL is a list of well distributed numbers which total addition is NB
        LN = NB
        NU = len(trim)//2
        NL = random_distribution(NU,NB)
        while not NL == []:
            captured_list = the_list[trim[0]:trim[1]+1]
            LN -= NL[0]
            blanked_list = random_long_blank(captured_list,NL[0],LB,exception)
            the_list[trim[0]:trim[1]+1] = blanked_list
            
            trim.pop(0)
            trim.pop(0)
            NL.pop(0)
            
        return the_list

################################################################################
################################################################################
################################################################################

class Problem_Precedent(object):

    def __init__(self):

        self.Whole_text = []
        self.Problem_text = ""
        self.Title = ""
        self.dic = {}
        
    def blank(self,NB,LB,keyword=None):

        Problem_text_list = random_point_blank(self.Whole_text,NB,LB)
        self.Problem_text = " ".join(Problem_text_list) +"\n"
        return self.Problem_text

    def Googling(self,keyword_list):
        num = 0 
        while keyword_list[num] in self.Whole_text:
            num += 1
            if num == len(keyword_list):
                return True
        return False


            
################################################################################
################################################################################
################################################################################
        

class Problem_Provision(object):
    
    def __init__(self):

        self.Whole_text = []
        self.Problem_text = ""
        self.Last_class_name = ["","","",""]
        self.Last_class = 0
        self.Keyword = []

    def put_data(self,Line_Type,Text):
        """
        데이터 입력함수. 
        조, 항, 호, 목의 본문은 각각 하나의 Line에 할당된다.
        즉, Line_Type은 "제5조의2", "제1항" 같은 Str 타입의 정보를 담고 있다.
        다만, 하나의 Line에 조,항이 같이 포함되는 등, 복잡한 경우가 있어 주의해야한다.
        """ 
        (val_A, val_P, val_S, val_I) = Line_Type

        # 문제 class에서는 조문을 구분하는 가장 높은 단위는 '조'다. 
        if val_A:
            self.Last_class_name[0]= val_A

        # 제1항 다음에 나오는 본문이 제2항이라면, present class는 1, last class도 1. 
        # 제1호 다목 다음에 나오는 본문이 2호라면, present class는 2, last class는 3. 
        for i in range(len(Line_Type)):
            if Line_Type[i]:
                present_class = i

        #제1항제1호다목 -> 제1항제2호 ;
        #["제10조","제1항","제1호",다목] -> ["제10조","제1항","제2호",""]
        self.Last_class_name[present_class] = Line_Type[present_class]
        if present_class <= self.Last_class:
            for i in range(len(self.Last_class_name)):
                if i > present_class:
                    self.Last_class_name[i] = ""

        total_class_name = ""
        for name in self.Last_class_name:
            total_class_name += str(name)
        self.Whole_text.append([total_class_name,Text])        
        self.Last_class = present_class

    
    def blank(self,NB,LB):
        """
        한 조항의 모든 항,호,목의 텍스트를 일률적으로 합친다.
        모드에 따라 빈칸을 내는 방법이 다르나, 재조립하는 방법은 같다.
        """
        self.Problem_text = All_Text(self.Whole_text)
        Problem_text_list1 = [word.strip(" ") for word in self.Problem_text.split(" ")]
        # 조문에 LB만큼 연속된 빈칸을 NB개 랜덤하게 만든다. 
        Problem_text_list2 = random_point_blank(Problem_text_list1,NB,LB)
        #빈칸을 넣은 텍스트를 재조립  
        self.Problem_text = " ".join(Problem_text_list2)
        return self.Problem_text

    def check_deleted(self):
        data = self.Whole_text[0][1]
        data1 = data.split(" ")
        data2 = [word.strip(" ") for word in data1 if word.strip(" ")]
        if data2[1] == "삭제":
            return True
        else :
            return False
        
    def Googling(self,keyword_list):
        Googling_text = All_Text(self.Whole_text)
        num = 0
        while keyword_list[num] in Googling_text:
            num += 1
        if num == len(keyword_list):
            return True
        else :
            return False


################################################################################
################################################################################
################################################################################


def Article_finder(text_list):
    # ['제3조',...]와 같이 구분된 텍스트에서 '제'를 찾고, '제3조의2' 같은 신설조항을 구분해준다. 
    text = text_list[0]
    total = ""
    if text[0] == '제':
        for letter in text:
            try :
                int(letter)
                total += letter 
            except :
                if letter == OrderInLaw[4] or letter == "의":
                    total += letter
                elif letter in OrderInLaw[0:4]:
                    return 0
        while total[-1] == "의":
            total = replaceRight(total,'의','',1)
        return "제" + total 
    else :
        return 0 


def Paragraph_finder(text_list):
    # 텍스트에서 처음 표기된 특수문자를 찾고 값을 되돌린다.(예: '①','②'...)
    total = ""
    for unit in text_list:
        for number in paragraph_symbol:
            if number in unit:
                total = number
                break 
            else :
                continue
    if total == "":
        return 0
    else :
        return '제' + total + '항'
    

def Subparagraph_finder(text_list):
    # '1.' 같이 적힌 호 표기를 찾는다. 온점을 찾은 다음, int 값으로 변환가능한지 본다.  
    total = ""
    text = text_list[0]
    total = text[:-1]
    if text[-1] == "." :
        try :
            int(text[:-1])
        except :
            return 0
        return "제" + total + "호"
    else :
        return 0 

def Item_finder(text_list):
    # '가.' 같이 적힌 목 표기를 찾는다. 온점을 찾은 다음, 가나다라...중 하나인지 확인한다. 
    total = ""
    text = text_list[0]
    total = text[:-1]
    if text[-1] == "." and (total in item_symbol):
        return total + "목"
    else :
        return 0 


################################################################################
################################################################################
################################################################################



def data_acquisition(templit):
    Problem_list = []
    Problem_unit = None
    switch = None
    Line_Type = None 
    
    for row, paragraph in enumerate(templit.paragraphs):

        line_test = paragraph.text.strip(' ').strip('\n')
        line = [word.strip(' ') for word in paragraph.text.split(' ') if word.strip(' ')]
        if line_test == "" or line == []:
            continue
        
        #Switch : Precedent or Provision / Normal Context or Title 
        if line[0][0] == "<":
            if Problem_unit:
                Problem_list.append(Problem_unit)
            Problem_unit = Problem_Precedent()
            Problem_unit.Title = line[0]
            switch = "Precedent"
        elif Article_finder(line):
            if Problem_unit and switch == "Precedent" :
                Problem_list.append(Problem_unit)
            elif Problem_unit and switch == "Provision" :
                if not Problem_unit.check_deleted():
                    Problem_list.append(Problem_unit)
            Problem_unit = Problem_Provision()
            switch = "Provision"
        # Data Type 
        if switch == "Precedent":

            Problem_unit.Whole_text += line
        elif switch == "Provision":
            Line_Type = (Article_finder(line), Paragraph_finder(line), Subparagraph_finder(line), Item_finder(line))
            if Line_Type == (0,0,0,0) or Line_Type == ("","","",""):
                continue
            Arranged_text = " ".join(line) + "\n" + " "
            Problem_unit.put_data(Line_Type,Arranged_text)
        else :
            continue 

    Problem_list.append(Problem_unit)
    return Problem_list

def Respawn_Problem(Vacant,Problem_list,NP,NB,LB,keyword_list):

    if not keyword_list == []:
        for problem in Problem_list:
            if problem.googling(keyword_list) :
                Problem_list.remove(problem)

    for i in range(NP):
        if Problem_list == []:
            break
        new_problem = random.choice(Problem_list)
        Problem_list.remove(new_problem)
        new_problem.blank(NB,LB)
        Vacant.add_paragraph(new_problem.Problem_text)
        
    Vacant.save(Result)

################################################################################
################################################################################
################################################################################

Problem_list = data_acquisition(templit)        
Respawn_Problem(Vacant,Problem_list,NP,NB,LB,keyword_list)      










































            
