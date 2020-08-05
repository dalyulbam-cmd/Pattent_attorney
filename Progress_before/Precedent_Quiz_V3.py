import os, random, itertools, time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

################################################################################

common = ["아니다.","따라","또는","이는","판시하였다.","있다.","것","수","여기서","여하","이","것은",
          "뿐만","의하여","없다.","각","이를","이러한","그","대한","등","것인","할","및","라고","수"
          "것이고","것으로서","있어서","있어","것이고,","하여","볼"]

################################################################################
"""
판례집 리스트 

    특허
        특허심판판례.docx
        특허정의판례.docx
        특허요건판례.docx
    상표
    민법
        민법총칙법인판례.docx
"""

# NP = Number of Problem, NB = Number of Blank, LB = Length of Blank
# 문제 수를 조절하세요. 
NP,NB,LB = 6,5,1

# 원하는 법조문 Templit을 가져와 문제를 만드세요. 
Current = os.getcwd()
print(Current)
templit = Document('자료/특허요건판례.docx')
Vacant = Document('Today.docx')
Result = "Today's_Problem_" + str(time.strftime("%y-%m-%d-%H-%M")) +".docx"

Keyword = None

################################################################################

class Problem(object):

    def __init__(self):

        self.Whole_text = []
        self.Problem_text = ""
        self.Title = ""
        self.dic = {}
        
    def blank(self,NB,LB,mode=1,keyword=None):

        self.internal_statistic()
        if mode == 0:
            Problem_text_list = random_blank(self.Whole_text,NB,LB)
        elif mode == 1:
            Problem_text_list = random_long_blank(self.Whole_text,NB,LB)
        elif mode == 2:
            Problem_text_list = random_point_blank(self.Whole_text,NB,LB)
        self.Problem_text = " ".join(Problem_text_list) +"\n"
        return self.Problem_text

    def internal_statistic(self):

        Target = self.Whole_text
        dic = {}
        for word in Target:
            if word in dic.keys():
                dic[word] += 1
            else :
                dic[word] = 1
        self.dic = dic 
        size = sum(dic.values())
        
        
################################################################################

def random_sublist(the_list,NB,LB):
    order_list = []
    sublist = []
    max_index = len(the_list)-1
    
    def drawing():
        n = random.randint(0,max_index)
        if n in order_list:
            return drawing()
        else :
            order_list.append(n)
            return order_list
    for i in range(NB):
        drawing()
    for n in order_list:
        sublist.append(the_list[n])
    return order_list, sublist

def random_long_sublist(the_list,NB,LB):
    order_list = []
    sublist = []
    end = len(the_list)-1 
    start = LB -1 
    
    def drawing(order_list):
        n = random.randint(start,end)
        group = []
        for i in range(LB):
            value = n-i
            if value in order_list:
                return drawing(order_list)
            group.append(n-i)
        order_list += group
        return order_list
    for i in range(NB):
        order_list = drawing(order_list)
    for j in order_list:
        sublist.append(the_list[j])
    return order_list, sublist

def random_long_blank(the_list,NB,LB):
    order_list, sublist = random_long_sublist(the_list,NB,LB)
    for i in range(len(sublist)):
        while sublist[i] in common:
            little_order_list, little_sublist = random_sublist(the_list,1,1)
            if little_order_list[0] in order_list:
                continue
            order_list[i] = little_order_list[0]
            sublist[i] = little_sublist[0]
        word = sublist[i]
        blank = ""
        for j in range(len(word)):
            #글자 수를 알 수 없게 하는 장치. 
            if j%4 ==0 : 
                blank += "    "
        blank = "(" + blank + ")"
        sublist[i] = blank
        
    for i in range(len(sublist)):
        the_list[order_list[i]] = sublist[i]
    return the_list    

def random_point_blank(the_list,NB,LB):
    trim = []
    for word in the_list:
        if '"' in word :
            trim.append(the_list.index(word))
    if trim == []:
        return random_long_blank(the_list,NB,LB)
    else :
        edited_list = the_list[trim[0]:trim[1]+1]
        return random_long_blank(edited_list,NB,LB)
            

################################################################################

def data_acquisition(templit):
    Problem_list = []
    Problem_unit = None

    for row, paragraph in enumerate(templit.paragraphs):
        line = [word.split('\n') for word in paragraph.text.split(' ') if word.strip(' ')]
        for i in range(len(line)):
            if type(line[i]) == list:
                line[i] = "\n ".join(line[i]).split(" ")
        line = list(itertools.chain(*line))
        line = [word.strip(" ") for word in line if word.strip(" ")]
        if line == []:
            continue

        if line[0][0] == "<":
            if Problem_unit:
                Problem_list.append(Problem_unit)
            Problem_unit = Problem()
            Problem_unit.Title = line[0]
        else :
            Problem_unit.Whole_text += line

    Problem_list.append(Problem_unit)
    return Problem_list

def Respawn_Problem(Vacant,Problem_list,NP,NB,LB):

    for i in range(NP):
        new_problem = random.choice(Problem_list)
        Problem_list.remove(new_problem)
        new_problem.blank(NB,LB,mode=2,keyword=None)
        Vacant.add_paragraph(new_problem.Problem_text)
    Vacant.save(Result)

################################################################################
    
Problem_list = data_acquisition(templit)        
Respawn_Problem(Vacant,Problem_list,NP,NB,LB)                
        
        
        
